#!/usr/bin/env python3
"""Render the Noesis patent submission set to .docx -- Arial 11, 1.5 spacing.

Generates clean upload documents from the markdown sources, stripping internal
draft/strategy notes (leading blockquotes, "(DRAFT -- PRIVATE)" tags, and any
"## Notes before submission" tail) so only the filing-ready text remains.

Outputs to Desktop:
  - Noesis-US-FFL-Petition.docx   (the 37 CFR 5.13 petition)
  - Noesis-SB15A-micro-entity.docx
  - Noesis-UKIPO-description.docx  (the material filed abroad)
  - Noesis-Provisional-Spec.docx   (full provisional, for the attorney's read)
"""
import os
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

DESK = os.path.expanduser("~/Desktop")


def clean(text: str) -> str:
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)            # bold
    text = re.sub(r"(?<!\*)\*(?!\*)(.+?)\*(?!\*)", r"\1", text)  # italic
    text = re.sub(r"\[(.+?)\]\((.+?)\)", r"\1", text)        # links -> text
    text = text.replace("`", "")
    text = re.sub(r"\(DRAFT.*?PRIVATE\)", "", text, flags=re.I)
    return text.strip()


def style(doc: Document) -> None:
    n = doc.styles["Normal"]
    n.font.name = "Arial"
    n.font.size = Pt(11)
    n.paragraph_format.line_spacing = 1.5
    n.paragraph_format.space_after = Pt(6)
    for h, sz in (("Heading 1", 13), ("Heading 2", 11.5)):
        st = doc.styles[h]
        st.font.name = "Arial"
        st.font.size = Pt(sz)
        st.font.bold = True
        st.font.color.rgb = RGBColor(0, 0, 0)
        st.paragraph_format.line_spacing = 1.5


def render(src: str, out: str, start: str | None = None, stop: str | None = None,
           title_marker: str | None = None) -> None:
    lines = open(src, encoding="utf-8").read().splitlines()
    # drop everything before `start` (a heading line) if given
    if start:
        for i, l in enumerate(lines):
            if l.strip() == start:
                lines = lines[i:]
                break
    # drop everything from `stop` onward (an internal tail section)
    if stop:
        for i, l in enumerate(lines):
            if l.strip() == stop:
                lines = lines[:i]
                break
    # drop leading blockquote note + leading blank/--- lines
    while lines and (lines[0].lstrip().startswith(">") or not lines[0].strip()
                     or lines[0].strip() == "---"):
        lines.pop(0)

    doc = Document()
    style(doc)
    for raw in lines:
        line = raw.rstrip()
        s = line.strip()
        if not s or s == "---":
            continue
        if title_marker and s == title_marker:
            continue
        if s.startswith("|"):  # markdown table row -> tab-joined paragraph
            cells = [c.strip() for c in s.strip("|").split("|")]
            if all(set(c) <= set("-: ") for c in cells):
                continue  # separator row
            doc.add_paragraph("    ".join(cells))
            continue
        if s.startswith("#### "):
            doc.add_heading(clean(s[5:]), level=2)
        elif s.startswith("### "):
            doc.add_heading(clean(s[4:]), level=2)
        elif s.startswith("## "):
            doc.add_heading(clean(s[3:]), level=1)
        elif s.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(clean(s[2:]))
            r.font.name = "Arial"; r.font.size = Pt(14); r.font.bold = True
            p.paragraph_format.line_spacing = 1.5
        elif re.match(r"^[-*] ", s):
            doc.add_paragraph(clean(s[2:]), style="List Bullet")
        elif re.match(r"^\d+\. ", s):
            doc.add_paragraph(clean(re.sub(r"^\d+\.\s*", "", s)), style="List Number")
        else:
            doc.add_paragraph(clean(s))
    doc.save(out)
    print(f"wrote {out}")


render("US-FFL-PETITION.md", os.path.join(DESK, "Noesis-US-FFL-Petition.docx"),
       stop="## Notes before submission (honest, confirm at checkout)")
render("SB15A-micro-entity.md", os.path.join(DESK, "Noesis-SB15A-micro-entity.docx"))
render("UKIPO-FILING.md", os.path.join(DESK, "Noesis-UKIPO-description.docx"),
       start="## INVENTION TITLE", title_marker="## INVENTION TITLE")
render("PROVISIONAL-PoM-value-chain-v2.md", os.path.join(DESK, "Noesis-Provisional-Spec.docx"))
